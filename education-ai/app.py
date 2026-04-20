import streamlit as st
from google import genai
from streamlit_ace import st_ace
import time
import random
import json

# ── CONFIG ──────────────────────────────────────────────  
st.set_page_config(page_title="EduAgent AI", page_icon="🎓", layout="wide")

# ── SESSION STATE INIT ───────────────────────────────────
if "show_home_output" not in st.session_state:
    st.session_state.show_home_output = False
if "page" not in st.session_state:
    st.session_state.page = "home"
if "tech_stage" not in st.session_state:
    st.session_state.tech_stage = "setup"
if "home_topic" not in st.session_state:
    st.session_state.home_topic = ""
if "home_output" not in st.session_state:
    st.session_state.home_output = ""
if "home_research" not in st.session_state:
    st.session_state.home_research = ""     # setup | interview | report
if "tech_questions" not in st.session_state:
    st.session_state.tech_questions = []
if "tech_answers" not in st.session_state:
    st.session_state.tech_answers = {}
if "tech_q_index" not in st.session_state:
    st.session_state.tech_q_index = 0
if "tech_start_time" not in st.session_state:
    st.session_state.tech_start_time = None
if "tech_total_seconds" not in st.session_state:
    st.session_state.tech_total_seconds = 0
if "tech_config" not in st.session_state:
    st.session_state.tech_config = {}
if "tech_report" not in st.session_state:
    st.session_state.tech_report = ""
if "hr_stage" not in st.session_state:
    st.session_state.hr_stage = "setup"
if "hr_config" not in st.session_state:
    st.session_state.hr_config = {}
if "hr_chat" not in st.session_state:
    st.session_state.hr_chat = []
if "hr_turn" not in st.session_state:
    st.session_state.hr_turn = 0
if "hr_start_time" not in st.session_state:
    st.session_state.hr_start_time = None
if "hr_report" not in st.session_state:
    st.session_state.hr_report = ""

# ── GEMINI CLIENT ────────────────────────────────────────
client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
MODEL_ID = "gemini-3.1-flash-lite-preview"

# ── STYLES ──────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@300;400;500;600;700&family=DM+Mono:wght@300;400;500&display=swap');

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"] {
    background-color: #060d10 !important;
    color: #e0eef2 !important;
    font-family: 'Sora', sans-serif !important;
}
/* ── Fake Fullscreen Loader ── */
.edu-loader {
    position: fixed;
    inset: 0;
    background: rgba(6, 13, 16, 0.92);
    backdrop-filter: blur(8px);
    z-index: 99999;

    display: flex;
    align-items: center;
    justify-content: center;
    flex-direction: column;

    animation: fadeIn 0.25s ease;
}

.edu-loader-spinner {
    width: 52px;
    height: 52px;
    border: 3px solid rgba(0,210,180,0.15);
    border-top: 3px solid #00d4b4;
    border-radius: 50%;
    animation: spin 0.9s linear infinite;
    margin-bottom: 1rem;
}

.edu-loader-text {
    font-size: 0.75rem;
    color: #8eb8c0;
    font-family: 'DM Mono', monospace;
    letter-spacing: 0.12em;
}

@keyframes spin {
    to { transform: rotate(360deg); }
}

@keyframes fadeIn {
    from { opacity: 0 }
    to { opacity: 1 }
}

#MainMenu, footer, header { visibility: hidden; }
[data-testid="stToolbar"] { display: none; }
[data-testid="stDecoration"] { display: none; }
[data-testid="stSidebarNav"] { display: none; }
[data-testid="stSidebar"] { display: none; }

[data-testid="stMain"] { background: transparent !important; }
.block-container { padding: 0 !important; max-width: 100% !important; }

/* ── Navbar ── */
.edu-navbar {
    position: sticky; top: 0; z-index: 999;
    display: flex; align-items: center; justify-content: space-between;
    padding: 1rem 3rem;
    background: rgba(6,13,16,0.92);
    backdrop-filter: blur(24px);
    border-bottom: 1px solid rgba(0,210,180,0.08);
}
.edu-navbar-logo {
    display: flex; align-items: center; gap: 0.6rem;
    font-size: 1.1rem; font-weight: 600; color: #e0eef2;
}
.edu-navbar-logo-dot {
    width: 8px; height: 8px; border-radius: 50%;
    background: #00d4b4; box-shadow: 0 0 10px #00d4b4;
    animation: pulse-dot 2.4s ease-in-out infinite;
}
@keyframes pulse-dot { 0%,100%{transform:scale(1);opacity:1;} 50%{transform:scale(1.4);opacity:0.7;} }
.edu-navbar-badge {
    font-size: 0.72rem; padding: 0.35rem 0.9rem; border-radius: 20px;
    border: 1px solid rgba(0,212,180,0.3); color: #00d4b4;
    font-family: 'DM Mono', monospace; background: rgba(0,212,180,0.06);
}



/* ── Hero ── */
.edu-hero {
    position: relative; min-height: 56vh;
    display: flex; align-items: center;
    padding: 2rem 3rem; overflow: hidden;
}
.edu-orb {
    position: absolute; right: 4%; top: 50%; transform: translateY(-50%);
    width: 36vw; height: 36vw; max-width: 460px; max-height: 460px;
    border-radius: 50%;
    background: radial-gradient(circle at 38% 42%, rgba(0,210,180,0.22) 0%, rgba(0,140,120,0.14) 30%, rgba(0,60,80,0.08) 60%, transparent 75%);
    box-shadow: inset 0 0 80px rgba(0,210,180,0.08), 0 0 120px rgba(0,210,180,0.06);
    animation: orb-breathe 6s ease-in-out infinite; pointer-events: none;
}
.edu-orb::before { content:''; position:absolute; inset:0; border-radius:50%; border:1px solid rgba(0,210,180,0.12); animation:orb-ring 6s ease-in-out infinite; }
.edu-orb::after  { content:''; position:absolute; inset:8%; border-radius:50%; border:1px solid rgba(0,210,180,0.06); }
.edu-orb-particles { position:absolute; inset:0; border-radius:50%; overflow:hidden; }
.edu-particle { position:absolute; width:2px; height:2px; background:#00d4b4; border-radius:50%; opacity:0; animation:particle-fade var(--dur,3s) ease-in-out var(--delay,0s) infinite; }
@keyframes particle-fade { 0%,100%{opacity:0;transform:scale(0.5);} 50%{opacity:var(--op,0.6);transform:scale(1);} }
@keyframes orb-breathe { 0%,100%{transform:translateY(-50%) scale(1);} 50%{transform:translateY(-50%) scale(1.04);} }
@keyframes orb-ring { 0%,100%{opacity:0.5;} 50%{opacity:1;} }
.edu-glow-tl { position:absolute; top:-10%; left:-5%; width:40vw; height:40vw; border-radius:50%; background:radial-gradient(circle,rgba(0,100,120,0.06) 0%,transparent 70%); pointer-events:none; }
.edu-glow-br { position:absolute; bottom:-15%; right:20%; width:30vw; height:30vw; border-radius:50%; background:radial-gradient(circle,rgba(0,60,80,0.07) 0%,transparent 70%); pointer-events:none; }
.edu-hero-content { position:relative; z-index:2; max-width:52%; }
.edu-hero-eyebrow { font-size:0.72rem; letter-spacing:0.2em; text-transform:uppercase; color:#00d4b4; margin-bottom:1.4rem; font-family:'DM Mono',monospace; }
.edu-hero-title { font-size:clamp(2.4rem,4.5vw,4.2rem); font-weight:700; line-height:1.08; color:#e8f4f6; letter-spacing:-0.02em; margin-bottom:1.4rem; }
.edu-hero-title .accent { color:transparent; -webkit-text-stroke:1px rgba(0,210,180,0.5); }
.edu-hero-desc { font-size:0.92rem; line-height:1.8; color:#4a7a86; max-width:440px; margin-bottom:2.8rem; font-weight:300; }
.edu-hero-floats { position:absolute; right:3%; top:18%; z-index:3; display:flex; flex-direction:column; gap:1rem; }
.edu-float-pill { background:rgba(8,22,26,0.82); border:1px solid rgba(0,210,180,0.15); border-radius:10px; padding:0.7rem 1.1rem; backdrop-filter:blur(16px); text-align:right; }
.edu-float-pill-title { font-size:0.78rem; font-weight:600; color:#b8d8df; }
.edu-float-pill-sub { font-size:0.66rem; color:#3e6470; margin-top:2px; font-family:'DM Mono',monospace; }
.edu-feature-cards { position:absolute; bottom:6%; right:4%; display:flex; gap:0.8rem; z-index:3; }
.edu-feature-card { background:rgba(8,22,26,0.85); border:1px solid rgba(0,210,180,0.12); border-radius:12px; padding:0.9rem 1.2rem; backdrop-filter:blur(20px); min-width:140px; transition:border-color 0.3s; }
.edu-feature-card:hover { border-color:rgba(0,210,180,0.3); }
.edu-feature-card-icon { font-size:1.2rem; margin-bottom:0.4rem; display:block; }
.edu-feature-card-title { font-size:0.78rem; font-weight:600; color:#b8d8df; display:block; }
.edu-feature-card-sub { font-size:0.65rem; color:#2e5058; font-family:'DM Mono',monospace; margin-top:2px; display:block; }

/* ── Input & Form ── */
[data-testid="stTextInput"] input {
    background:rgba(10,26,30,0.9) !important; border:1px solid rgba(0,210,180,0.2) !important;
    border-radius:12px !important; color:#e0eef2 !important; font-family:'Sora',sans-serif !important;
    font-size:1rem !important; padding:1rem 1.4rem !important; height:auto !important;
}
[data-testid="stTextInput"] input:focus { border-color:rgba(0,210,180,0.5) !important; box-shadow:0 0 0 3px rgba(0,210,180,0.08) !important; outline:none !important; }
[data-testid="stTextInput"] input::placeholder { color:#2e5058 !important; }
[data-testid="stTextInput"] label { display:none !important; }
[data-testid="stForm"] { border:none !important; padding:0 !important; background:transparent !important; }
[data-testid="stFormSubmitButton"] button {
    background:linear-gradient(135deg,rgba(0,180,160,0.15) 0%,rgba(0,120,110,0.1) 100%) !important;
    border:1px solid rgba(0,210,180,0.35) !important; border-radius:10px !important;
    color:#00d4b4 !important; font-family:'Sora',sans-serif !important;
    font-size:0.88rem !important; font-weight:600 !important; letter-spacing:0.06em !important;
    padding:0.1rem 1.2rem !important; min-width:100px !important; white-space:nowrap !important;
    height:auto !important; width:100% !important; text-transform:uppercase !important; cursor:pointer !important;
}
[data-testid="stFormSubmitButton"] button:hover {
    background:linear-gradient(135deg,rgba(0,210,180,0.22) 0%,rgba(0,160,140,0.16) 100%) !important;
    border-color:rgba(0,210,180,0.65) !important; box-shadow:0 0 24px rgba(0,210,180,0.15) !important;
    transform:translateY(-1px) !important;
}

/* ── General Buttons ── */
.stButton > button {
     background:linear-gradient(135deg,rgba(0,210,180,0.22) 0%,rgba(0,160,140,0.16) 100%) !important;
    border:1px solid rgba(0,210,180,0.35) !important; border-radius:10px !important;
    color:#00d4b4 !important; font-family:'Sora',sans-serif !important;
    font-size:0.88rem !important; font-weight:600 !important; letter-spacing:0.06em !important;
    padding:0, 1.5rem !important; height:auto !important; width:auto !important;
  text-transform: none !important; letter-spacing: 0 !important; cursor:pointer !important; transition:all 0.25s ease !important;
}
.stButton > button:hover {
    background:linear-gradient(135deg,rgba(0,210,180,0.22) 0%,rgba(0,160,140,0.16) 100%) !important;
    border-color:rgba(0,210,180,0.65) !important; box-shadow:0 0 24px rgba(0,210,180,0.15) !important;
    transform:translateY(-1px) !important;
}
.stButton > button:active { transform:translateY(0px) !important; }
[data-testid="stBaseButton-secondary"]:has(+ div [data-testid="stBaseButton-secondary"]) { display:none; }
.feat-card { background:rgba(8,22,26,0.85); border:1px solid rgba(0,210,180,0.15); border-top:2px solid #00d4b4; border-bottom:none; border-radius:14px 14px 0 0; padding:1.6rem 1.4rem 0.8rem; text-align:center; }
.feat-btn .stButton > button { background:rgba(8,22,26,0.85) !important; border:1px solid rgba(0,210,180,0.15) !important; border-top:none !important; border-radius:0 0 14px 14px !important; color:#00d4b4 !important; font-size:0.8rem !important; padding:0.6rem !important; text-transform:none !important; letter-spacing:0.04em !important; width:60% !important; margin-top:0 !important; }
.feat-btn .stButton > button:hover { background:rgba(0,210,180,0.08) !important; border-color:rgba(0,210,180,0.35) !important; transform:none !important; box-shadow:none !important; }
.feat-btn .stButton { width:100% !important; }
/* ── Selectbox smaller + margin ── */
[data-testid="stSelectbox"] { margin-bottom: 0.5rem !important; }
[data-testid="stSelectbox"] > div > div {
    background:rgba(10,26,30,0.9) !important; border:1px solid rgba(0,210,180,0.2) !important;
    border-radius:10px !important; color:#e0eef2 !important; font-family:'Sora',sans-serif !important;
font-size:0.85rem !important; padding: 0.55rem 1rem !important; height:auto !important;
}
[data-testid="stSelectbox"] label {
    color:#8eb8c0 !important; font-family:'DM Mono',monospace !important;
    font-size:0.72rem !important; letter-spacing:0.08em !important; text-transform:uppercase !important;
    margin-bottom: 0.3rem !important;

}

/* ── Technical interview layout ── */
.tech-setup-card {
    background: rgba(8,22,26,0.85);
    border: 1px solid rgba(0,210,180,0.12);
    border-radius: 16px;
    padding: 2.5rem;
    max-width: 560px;
    margin: 0rem ;
}
.tech-setup-title {
    font-size: 1.8rem; font-weight: 700; color: #e8f4f6;
    margin-bottom: 0.5rem; letter-spacing: -0.02em;
}
.tech-setup-sub {
    font-size: 0.85rem; color: #4a7a86; margin-bottom: 2rem;
    font-family: 'DM Mono', monospace;
}
.tech-question-panel {
    background: rgba(8,22,26,0.85);
    border: 1px solid rgba(0,210,180,0.12);
    border-radius: 12px;
    padding: 1.5rem;
    height: 100%;
}
.tech-q-tag {
    font-size: 0.65rem; letter-spacing: 0.15em; text-transform: uppercase;
    color: #00d4b4; font-family: 'DM Mono', monospace; margin-bottom: 0.8rem;
}
.tech-q-title { font-size: 1.1rem; font-weight: 600; color: #e8f4f6; margin-bottom: 1rem; }
.tech-q-body { font-size: 0.88rem; line-height: 1.8; color: #4a7a86; margin-bottom: 1.2rem; }
.tech-q-section { font-size: 0.72rem; text-transform: uppercase; letter-spacing: 0.1em; color: #2e5058; font-family: 'DM Mono', monospace; margin-bottom: 0.4rem; margin-top: 1rem; }
.tech-q-example { background: rgba(0,210,180,0.04); border-left: 2px solid rgba(0,210,180,0.3); border-radius: 0 6px 6px 0; padding: 0.6rem 0.9rem; font-family: 'DM Mono', monospace; font-size: 0.8rem; color: #8eb8c0; margin-bottom: 0.5rem; }
.tech-timer-bar {
    background: rgba(8,22,26,0.9); border: 1px solid rgba(0,210,180,0.12);
    border-radius: 10px; padding: 0.8rem 1.5rem;
    display: flex; align-items: center; justify-content: space-between;
    margin-bottom: 1rem;
}
.tech-timer-label { font-size:0.7rem; text-transform:uppercase; letter-spacing:0.1em; color:#3e6470; font-family:'DM Mono',monospace; }
.tech-timer-value { font-size:1.2rem; font-weight:700; color:#00d4b4; font-family:'DM Mono',monospace; }
.tech-timer-value.warning { color: #f5a623; }
.tech-timer-value.danger  { color: #e24b4a; }

/* ── Report cards ── */
.report-card {
    background: rgba(8,22,26,0.85); border: 1px solid rgba(0,210,180,0.12);
    border-radius: 12px; padding: 1.2rem 1.5rem; margin-bottom: 1rem;
}
.report-pass  { border-left: 3px solid #02C39A; }
.report-partial { border-left: 3px solid #f5a623; }
.report-fail  { border-left: 3px solid #e24b4a; }
.report-badge {
    font-size: 0.65rem; font-weight: 600; text-transform: uppercase;
    letter-spacing: 0.1em; padding: 0.2rem 0.6rem; border-radius: 4px;
    font-family: 'DM Mono', monospace;
}
.badge-pass    { background: rgba(2,195,154,0.12); color: #02C39A; }
.badge-partial { background: rgba(245,166,35,0.12); color: #f5a623; }
.badge-fail    { background: rgba(226,75,74,0.12);  color: #e24b4a; }

/* ── Misc ── */
/* Page content margin */
[data-testid="stMain"] > div > div > div {
    padding-left: 3rem !important;
    padding-right: 3rem !important;
}
[data-testid="stMarkdownContainer"] h1,[data-testid="stMarkdownContainer"] h2,[data-testid="stMarkdownContainer"] h3 { color:#b8d8df !important; font-family:'Sora',sans-serif !important; font-weight:600 !important; margin-top:0rem !important; }
[data-testid="stMarkdownContainer"] p,[data-testid="stMarkdownContainer"] li { color:#4a7a86 !important; font-family:'Sora',sans-serif !important; font-size:0.9rem !important; line-height:1.8 !important; }
[data-testid="stMarkdownContainer"] strong { color:#8eb8c0 !important; }
[data-testid="stAlert"] { background:rgba(8,22,26,0.9) !important; border-left:3px solid #00d4b4 !important; border-radius:8px !important; font-family:'Sora',sans-serif !important; font-size:0.85rem !important; }
hr { border-color:rgba(0,210,180,0.08) !important; margin:0rem 0 !important; }
[data-testid="column"] { padding:0 0.4rem !important; }
::-webkit-scrollbar { width:4px; height:4px; }
::-webkit-scrollbar-track { background:#060d10; }
::-webkit-scrollbar-thumb { background:rgba(0,210,180,0.2); border-radius:2px; }
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════
# NAVBAR (always visible)
# ════════════════════════════════════════════════════════
st.markdown("""
<div class="edu-navbar">
    <div class="edu-navbar-logo">
        <div class="edu-navbar-logo-dot"></div>
        EduAgent
    </div>
    <div class="edu-navbar-badge">Agents · Live</div>
</div>
""", unsafe_allow_html=True)



# ════════════════════════════════════════════════════════
# AGENT HELPERS
# ════════════════════════════════════════════════════════
def generate_questions(role, difficulty, language, num_questions):
    prompt = f"""You are a technical interviewer. Generate exactly {num_questions} coding interview questions for a {role} position.
Difficulty: {difficulty}. The candidate will code in {language}.

Return ONLY a valid JSON array with no extra text. Each item must have exactly these keys:
- "title": short problem name
- "description": full problem statement (2-3 sentences)
- "examples": list of 2 dicts with "input" and "output" keys
- "constraints": list of 2-3 constraint strings
- "test_cases": list of 3 dicts with "input" and "expected" keys

Example format:
[{{"title":"Two Sum","description":"...","examples":[{{"input":"nums=[2,7,11,15], target=9","output":"[0,1]"}}],"constraints":["2 <= nums.length <= 10^4"],"test_cases":[{{"input":"[2,7,11,15], 9","expected":"[0,1]"}}]}}]"""
    response = client.models.generate_content(model=MODEL_ID, contents=prompt)
    text = response.text.strip()
    # Strip markdown fences if present
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]
    return json.loads(text.strip())

def evaluate_code(question, code, language):
    prompt = f"""You are a strict code reviewer. Evaluate this {language} code for the following problem.

Problem: {question['title']}
Description: {question['description']}
Test cases: {json.dumps(question['test_cases'])}

Submitted code:
```{language}
{code}
```

Return ONLY a valid JSON object with these keys:
- "verdict": one of "Pass", "Partial", "Fail"
- "score": integer 0-100
- "test_results": list of dicts with "input", "expected", "verdict" ("Pass"/"Fail")
- "feedback": 2-3 sentences on code quality and correctness
- "time_complexity": e.g. "O(n)"
- "space_complexity": e.g. "O(1)"
"""
    response = client.models.generate_content(model=MODEL_ID, contents=prompt)
    text = response.text.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]
    return json.loads(text.strip())

def generate_final_report(config, answers, evaluations):
    total_score = round(sum(e.get("score", 0) for e in evaluations) / len(evaluations)) if evaluations else 0
    prompt = f"""You are a senior technical interviewer. Write a final interview report.

Candidate info: Role={config['role']}, Difficulty={config['difficulty']}, Language={config['language']}
Overall score: {total_score}/100
Questions attempted: {len(answers)}/{config['num_questions']}

Per-question scores: {[e.get('score',0) for e in evaluations]}
Per-question verdicts: {[e.get('verdict','') for e in evaluations]}

Write a concise report with these sections (use markdown):
## Overall Performance
## Strong Areas
## Weak Areas  
## Topics to Revise
## Recommendation (Hire / Consider / Reject)
"""
    response = client.models.generate_content(model=MODEL_ID, contents=prompt)
    return response.text

def researcher_agent(topic):
    prompt = f"You are a Researcher. Provide deep scientific facts about {topic}."
    response = client.models.generate_content(model=MODEL_ID, contents=prompt)
    return response.text

def writer_agent(raw_data):
    prompt = f"You are a Writer. Turn this into a professional study guide with structured headers and bullet points:\n{raw_data}"
    response = client.models.generate_content(model=MODEL_ID, contents=prompt)
    return response.text

def get_timer_seconds(num_questions, difficulty):
    table = {
        "Easy":   {5: 25*60,  10: 45*60,  15: 60*60},
        "Medium": {5: 40*60,  10: 75*60,  15: 110*60},
        "Hard":   {5: 60*60,  10: 120*60, 15: 180*60},
        "Mixed":  {5: 45*60,  10: 90*60,  15: 120*60},
    }
    return table.get(difficulty, {}).get(num_questions, 60*60)

def format_time(seconds):
    m, s = divmod(max(0, int(seconds)), 60)
    h, m = divmod(m, 60)
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"

def fake_loader(text="Processing...", duration=1.2):
    placeholder = st.empty()

    placeholder.markdown(f"""
    <div class="edu-loader">
        <div class="edu-loader-spinner"></div>
        <div class="edu-loader-text">{text}</div>
    </div>
    """, unsafe_allow_html=True)

    time.sleep(duration)

    placeholder.empty()
# ════════════════════════════════════════════════════════
# PAGE: HOME
# ════════════════════════════════════════════════════════
def show_home():
    particles_html = ""
    random.seed(42)
    for i in range(60):
        x = random.randint(5, 95)
        y = random.randint(5, 95)
        op = round(random.uniform(0.2, 0.8), 2)
        dur = round(random.uniform(2.0, 5.5), 1)
        delay = round(random.uniform(0, 4.0), 1)
        size = random.choice([1, 1, 2, 2, 2, 3])
        particles_html += f'<div class="edu-particle" style="left:{x}%;top:{y}%;width:{size}px;height:{size}px;--op:{op};--dur:{dur}s;--delay:{delay}s;"></div>'

    st.markdown(f"""
    <section class="edu-hero">
        <div class="edu-glow-tl"></div>
        <div class="edu-glow-br"></div>
        <div class="edu-orb"><div class="edu-orb-particles">{particles_html}</div></div>
        <div class="edu-hero-floats">
            <div class="edu-float-pill">
                <div class="edu-float-pill-title">Built with AI</div>
                <div class="edu-float-pill-sub">for maximum efficiency</div>
            </div>
            <div class="edu-float-pill">
                <div class="edu-float-pill-title">Smart Responses</div>
                <div class="edu-float-pill-sub">for every scenario</div>
            </div>
        </div>
        <div class="edu-hero-content">
            <div class="edu-hero-eyebrow">Multi-Agent Interview Platform</div>
            <h1 class="edu-hero-title">Ace Your Next<br><span class="accent">Interview</span></h1>
            <p class="edu-hero-desc">
                AI-powered technical and HR interviews with real-time feedback,
                code evaluation, and personalised reports — all in one place.
            </p>
        </div>
        <div class="edu-feature-cards">
            <div class="edu-feature-card">
                <span class="edu-feature-card-icon">⬡</span>
                <span class="edu-feature-card-title">Technical</span>
                <span class="edu-feature-card-sub">LeetCode-style</span>
            </div>
            <div class="edu-feature-card">
                <span class="edu-feature-card-icon">◈</span>
                <span class="edu-feature-card-title">HR Interview</span>
                <span class="edu-feature-card-sub">Behavioural</span>
            </div>
        </div>
    </section>
    """, unsafe_allow_html=True)

    # ── Topic search textbox (original functionality) ──
    st.markdown("<div style='padding:0 3rem 0;'>", unsafe_allow_html=True)
    with st.form(key="home_search_form"):
        col_l, col_m, col_btn, col_r = st.columns([1.3, 2, 1.1, 1])
        with col_m:
            topic = st.text_input("", placeholder="Enter a topic", label_visibility="collapsed")
        with col_btn:
            home_submit = st.form_submit_button("Submit")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Output appears below the textbox ──
    if home_submit and topic:
        st.session_state.show_home_output = True
        st.session_state.home_topic = topic

    if st.session_state.get("show_home_output"):
        topic = st.session_state.home_topic
        col_close, _ = st.columns([1, 8])
        with col_close:
            if st.button("✕ Close", key="close_output"):
                st.session_state.show_home_output = False
                st.rerun()
        if not st.session_state.get("home_output"):
            with st.status("⬡ Researcher Agent initializing...", expanded=True) as status:
                research_data = researcher_agent(topic)
                st.write("✦ Data harvested")
                time.sleep(0.6)
                final_output = writer_agent(research_data)
                status.update(label="✦ Complete", state="complete", expanded=False)
            st.session_state.home_output = final_output
            st.session_state.home_research = research_data
        tab1, tab2 = st.tabs(["✦  Study Guide", "◈  Raw Log"])
        with tab1:
            st.success(f"Study guide for: **{topic.upper()}**")
            st.text_area("Study Guide", st.session_state.home_output, height=380)
            st.download_button("↓ Download", st.session_state.home_output, file_name=f"{topic}_study_guide.md", mime="text/markdown")
        with tab2:
            st.text_area("Agent logs:", st.session_state.home_research, height=280)
            st.markdown("<div style='padding:1.5rem 3rem 0;'>", unsafe_allow_html=True)
            tab1, tab2 = st.tabs(["✦  Study Guide", "◈  Raw Log"])
            with tab1:
                st.success(f"Study guide for: **{topic.upper()}**")
                st.text_area("Study Guide", final_output, height=380)
                st.download_button("↓ Download", final_output, file_name=f"{topic}_study_guide.md", mime="text/markdown")
            with tab2:
                st.text_area("Agent logs:", research_data, height=280)
            st.markdown("</div>", unsafe_allow_html=True)

# ── Feature cards — Glassmorphism ──
    st.markdown("""
    <style>
    .feat-grid { display: flex; gap: 1.2rem; padding: 0 0rem; justify-content: center; }
    .feat-card-wrap { flex: 1; max-width: 380px; display: flex; flex-direction: column; }
    .feat-card2 {
        background: rgba(8,22,26,0.85); border: 1px solid rgba(0,210,180,0.15);
        border-top: 2px solid #00d4b4; border-radius: 14px 14px 0 0;
        padding: 1.6rem 1.4rem 0.8rem; text-align: center; flex: 1;
    }
    .feat-open-btn {
        display: block; width: 55%; margin: 0 auto;
        background: rgba(8,22,26,0.85); border: 1px solid rgba(0,212,180,0.3);
        border-top: none; border-radius: 0 0 14px 14px;
        color: #00d4b4; font-size: 0.8rem; padding: 0.6rem;
        text-align: center; cursor: pointer; font-family: 'DM Mono', monospace;
        letter-spacing: 0.04em; text-decoration: none;
        transition: background 0.2s ease;

       
    }
    .feat-open-btn:hover { background: rgba(0,210,180,0.08); border-color: rgba(0,210,180,0.35); color: #00d4b4; }
    </style>
    <div class="feat-grid">
        <div class="feat-card-wrap">
            <div class="feat-card2">
                <div style="font-size:2rem;margin-bottom:10px;">💻</div>
                <div style="font-size:1rem;font-weight:600;color:#e8f4f6;margin-bottom:6px;">Technical Interview</div>
                <div style="font-size:0.78rem;color:#4a7a86;line-height:1.6;margin-bottom:1rem;">Role-specific coding questions with AI evaluation and test case checking.</div>
            </div>
            <a class="feat-open-btn" href="?nav=technical"  target="_self">Open →</a>
        </div>
        <div class="feat-card-wrap">
            <div class="feat-card2">
                <div style="font-size:2rem;margin-bottom:10px;">🤝</div>
                <div style="font-size:1rem;font-weight:600;color:#e8f4f6;margin-bottom:6px;">HR Interview</div>
                <div style="font-size:0.78rem;color:#4a7a86;line-height:1.6;margin-bottom:1rem;">Behavioural and situational questions with communication scoring.</div>
            </div>
            <a class="feat-open-btn" href="?nav=hr" target="_self">Open →</a>
        </div>
        <div class="feat-card-wrap">
            <div class="feat-card2">
                <div style="font-size:2rem;margin-bottom:10px;">📄</div>
                <div style="font-size:1rem;font-weight:600;color:#e8f4f6;margin-bottom:6px;">Resume Generator</div>
                <div style="font-size:0.78rem;color:#4a7a86;line-height:1.6;margin-bottom:1rem;">Generate your own  professional resume by entering the detailes</div>
            </div>
            <a class="feat-open-btn" href="?nav=resume"  target="_self">Open →</a>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Handle navigation via query params
    params = st.query_params
    if "nav" in params:
        nav = params["nav"]
        st.query_params.clear()
        fake_loader("Loading...", 1.5)
        if nav == "technical":
            st.session_state.page = "technical"
            st.session_state.tech_stage = "setup"
        elif nav == "hr":
            st.session_state.page = "hr"
        elif nav == "resume":
            st.session_state.page = "resume"
        st.rerun()
# ════════════════════════════════════════════════════════
# PAGE: TECHNICAL — SETUP SCREEN
# ════════════════════════════════════════════════════════
def show_technical_setup():
    st.markdown("<div style='padding:0rem 3rem;'>", unsafe_allow_html=True)
    if st.button("← Back to Home", key="tech_back"):
        fake_loader("Going Back...", 1.5)
        st.session_state.page = "home"
        st.rerun()
    st.markdown("""
    <div style="margin:0rem 0rem 0rem 1.5rem;">
        <div style="font-size:0.7rem;letter-spacing:0.2em;text-transform:uppercase;color:#00d4b4;font-family:'DM Mono',monospace;margin-bottom:0rem,margin-left:3rem;">Technical Interview</div>
        <h2 style="font-size:2rem;font-weight:700;color:#e8f4f6;letter-spacing:-0.02em;">Configure Your Session</h2>
        <p style="font-size:0.88rem;color:#4a7a86;margin-top:0rem;">Choose your role and preferences to generate a personalised interview.</p>
    </div>
    """, unsafe_allow_html=True)

    with st.form("tech_setup_form"):
        st.markdown("<div style='max-width:520px;margin:0 0 0rem 0;'>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            role = st.selectbox("Target Role", [ "Select",
                "Python Developer", "Data Analyst", "ML Engineer",
                "Web Developer", "Data Engineer", "Java Developer",
                "Backend Developer", "Full Stack Developer"
            ])
            st.markdown("<div style='margin-top:0rem;'>", unsafe_allow_html=True)
            difficulty = st.selectbox("Difficulty", ["Select","Easy", "Medium", "Hard", "Mixed"])
            st.markdown("</div>", unsafe_allow_html=True)
        with col2:
            language = st.selectbox("Coding Language", ["Select","Python", "Java", "C++", "JavaScript"])
            st.markdown("<div style='margin-top:0rem;'>", unsafe_allow_html=True)
            num_questions = st.selectbox("Number of Questions", ["Select",5, 10, 15])
            st.markdown("</div>", unsafe_allow_html=True)
        col_btn, _ = st.columns([1, 2])  # control position

        with col_btn:
                    start = st.form_submit_button("⬡  Generate Interview")

    if start:
            fake_loader("Initializing...", 1.5)   
            try:
                if role == "Select" or difficulty == "Select" or language == "Select" or num_questions == "Select":
                    raise ValueError("Incomplete input")

                with st.spinner("Generating your personalised interview questions..."):
                    questions = generate_questions(role, difficulty, language, num_questions)

                st.session_state.tech_questions = questions
                st.session_state.tech_answers = {}
                st.session_state.tech_q_index = 0
                st.session_state.tech_start_time = time.time()
                st.session_state.tech_total_seconds = get_timer_seconds(num_questions, difficulty)

                st.session_state.tech_config = {
                    "role": role,
                    "difficulty": difficulty,
                    "language": language,
                    "num_questions": num_questions
                }

                st.session_state.tech_stage = "interview"
                st.rerun()

            except:
                st.warning("⚠️ Please complete all fields properly.")

    st.markdown("</div>", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════
# PAGE: TECHNICAL — INTERVIEW SCREEN
# ════════════════════════════════════════════════════════
def show_technical_interview():
    questions = st.session_state.tech_questions
    q_idx = st.session_state.tech_q_index
    config = st.session_state.tech_config
    total_secs = st.session_state.tech_total_seconds
    elapsed = time.time() - st.session_state.tech_start_time
    remaining = total_secs - elapsed

    # Auto-finish if timer expired
    if remaining <= 0:
        st.session_state.tech_stage = "report"
        st.rerun()

    st.markdown("<div style='padding:1rem 2rem;'>", unsafe_allow_html=True)

    # ── Timer bar ──
    timer_class = "danger" if remaining < 300 else ("warning" if remaining < 600 else "")
    st.markdown(f"""
    <div class="tech-timer-bar">
        <div>
            <div class="tech-timer-label">Time Remaining</div>
            <div class="tech-timer-value {timer_class}">{format_time(remaining)}</div>
        </div>
        <div style="text-align:right;">
            <div class="tech-timer-label">Question</div>
            <div class="tech-timer-value">{q_idx + 1} / {len(questions)}</div>
        </div>
        <div style="text-align:right;">
            <div class="tech-timer-label">Role · Difficulty</div>
            <div style="font-size:0.82rem;color:#4a7a86;font-family:'DM Mono',monospace;">{config['role']} · {config['difficulty']}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    q = questions[q_idx]

    # ── Two column layout ──
    left, right = st.columns([1, 1], gap="medium")

    with left:
        # Question panel
        examples_html = ""
        for ex in q.get("examples", []):
            examples_html += f'<div class="tech-q-example">Input: {ex["input"]}<br>Output: {ex["output"]}</div>'
        constraints_html = "".join(f"<li style='color:#4a7a86;font-size:0.82rem;'>{c}</li>" for c in q.get("constraints", []))

        st.markdown(f"""
        <div class="tech-question-panel">
            <div class="tech-q-tag">Question {q_idx + 1} · {config['difficulty']}</div>
            <div class="tech-q-title">{q['title']}</div>
            <div class="tech-q-body">{q['description']}</div>
            <div class="tech-q-section">Examples</div>
            {examples_html}
            <div class="tech-q-section">Constraints</div>
            <ul style="padding-left:1.2rem;margin-top:0.4rem;">{constraints_html}</ul>
        </div>
        """, unsafe_allow_html=True)

    with right:
        # Code editor
        st.markdown(f"""
        <div style="font-size:0.7rem;letter-spacing:0.12em;text-transform:uppercase;color:#3e6470;font-family:'DM Mono',monospace;margin-bottom:0.5rem;">
            {config['language']} · Code Editor
        </div>
        """, unsafe_allow_html=True)

        lang_map = {"Python": "python", "Java": "java", "C++": "c_cpp", "JavaScript": "javascript"}
        starter = {
            "Python": f"def solution():\n    # Write your solution here\n    pass\n",
            "Java": f"class Solution {{\n    public void solve() {{\n        // Write your solution here\n    }}\n}}\n",
            "C++": f"#include <bits/stdc++.h>\nusing namespace std;\n\nvoid solution() {{\n    // Write your solution here\n}}\n",
            "JavaScript": f"function solution() {{\n    // Write your solution here\n}}\n"
        }

        prev_code = st.session_state.tech_answers.get(q_idx, {}).get("code", starter[config['language']])

        code = st_ace(
            value=prev_code,
            language=lang_map[config['language']],
            theme="monokai",
            font_size=14,
            tab_size=4,
            height=340,
            key=f"ace_{q_idx}",
            auto_update=True
        )

        # Buttons row
        btn1, btn2, btn3 = st.columns([1, 1, 1])
        with btn1:
            if st.button("◀ Prev", key="prev_btn", disabled=(q_idx == 0)):
                fake_loader("Loading...", 1)
                st.session_state.tech_answers[q_idx] = {"code": code}
                st.session_state.tech_q_index -= 1
                st.rerun()
        with btn2:
            if st.button("Save", key="save_btn"):
                fake_loader("Saving...", 1.5)
                st.session_state.tech_answers[q_idx] = {"code": code}
                st.success("Saved!")
        with btn3:
            is_last = (q_idx == len(questions) - 1)
            label = "Finish ✓" if is_last else "Next ▶"
            if st.button(label, key="next_btn"):
                fake_loader("Loading...", 1.5)
                st.session_state.tech_answers[q_idx] = {"code": code}
                if is_last:
                    fake_loader("Generating...", 1.5)
                    st.session_state.tech_stage = "report"
                    st.rerun()
                else:
                    st.session_state.tech_q_index += 1
                    st.rerun()

    # Question nav dots
    dots_html = ""
    for i in range(len(questions)):
        answered = i in st.session_state.tech_answers
        active = i == q_idx
        color = "#00d4b4" if active else ("#02C39A" if answered else "#1e3a44")
        dots_html += f'<div style="width:10px;height:10px;border-radius:50%;background:{color};cursor:pointer;" title="Q{i+1}"></div>'

    st.markdown(f"""
    <div style="display:flex;gap:8px;align-items:center;justify-content:center;margin-top:1.2rem;flex-wrap:wrap;">
        {dots_html}
    </div>
    """, unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

    # Auto-refresh every 30 seconds to update timer
    time.sleep(0)
    st.markdown("""
    <script>
    setTimeout(function(){ window.location.reload(); }, 30000);
    </script>
    """, unsafe_allow_html=True)


# ════════════════════════════════════════════════════════
# PAGE: TECHNICAL — REPORT SCREEN
# ════════════════════════════════════════════════════════
def show_technical_report():
    st.markdown("<div style='padding:2rem 3rem;'>", unsafe_allow_html=True)
    st.markdown("""
    <div style="margin-bottom:2rem;">
        <div style="font-size:0.7rem;letter-spacing:0.2em;text-transform:uppercase;color:#00d4b4;font-family:'DM Mono',monospace;margin-bottom:0.5rem;">Interview Complete</div>
        <h2 style="font-size:2rem;font-weight:700;color:#e8f4f6;letter-spacing:-0.02em;">Your Results</h2>
    </div>
    """, unsafe_allow_html=True)

    questions = st.session_state.tech_questions
    answers = st.session_state.tech_answers
    config = st.session_state.tech_config

    # Evaluate all answered questions
    evaluations = []
    with st.status("⬡ Evaluating your submissions...", expanded=True) as status:
        for i, q in enumerate(questions):
            code = answers.get(i, {}).get("code", "")
            if code.strip():
                st.write(f"Evaluating Q{i+1}: {q['title']}...")
                try:
                    result = evaluate_code(q, code, config['language'])
                    evaluations.append(result)
                except:
                    evaluations.append({"verdict": "Fail", "score": 0, "feedback": "Could not evaluate.", "test_results": []})
            else:
                evaluations.append({"verdict": "Fail", "score": 0, "feedback": "No code submitted.", "test_results": []})
        status.update(label="✦ Evaluation complete", state="complete", expanded=False)

    # Overall score
    total_score = round(sum(e.get("score", 0) for e in evaluations) / len(evaluations)) if evaluations else 0
    time_used = time.time() - st.session_state.tech_start_time
    time_given = st.session_state.tech_total_seconds

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(f"""
        <div style="background:rgba(8,22,26,0.85);border:1px solid rgba(0,210,180,0.12);border-radius:12px;padding:1.5rem;text-align:center;">
            <div style="font-size:0.7rem;letter-spacing:0.1em;text-transform:uppercase;color:#3e6470;font-family:'DM Mono',monospace;">Overall Score</div>
            <div style="font-size:3rem;font-weight:700;color:#00d4b4;font-family:'DM Mono',monospace;">{total_score}</div>
            <div style="font-size:0.75rem;color:#3e6470;">out of 100</div>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown(f"""
        <div style="background:rgba(8,22,26,0.85);border:1px solid rgba(0,210,180,0.12);border-radius:12px;padding:1.5rem;text-align:center;">
            <div style="font-size:0.7rem;letter-spacing:0.1em;text-transform:uppercase;color:#3e6470;font-family:'DM Mono',monospace;">Time Used</div>
            <div style="font-size:2rem;font-weight:700;color:#b8d8df;font-family:'DM Mono',monospace;">{format_time(time_used)}</div>
            <div style="font-size:0.75rem;color:#3e6470;">of {format_time(time_given)}</div>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        attempted = sum(1 for i in range(len(questions)) if answers.get(i, {}).get("code", "").strip())
        st.markdown(f"""
        <div style="background:rgba(8,22,26,0.85);border:1px solid rgba(0,210,180,0.12);border-radius:12px;padding:1.5rem;text-align:center;">
            <div style="font-size:0.7rem;letter-spacing:0.1em;text-transform:uppercase;color:#3e6470;font-family:'DM Mono',monospace;">Attempted</div>
            <div style="font-size:2rem;font-weight:700;color:#b8d8df;font-family:'DM Mono',monospace;">{attempted}/{len(questions)}</div>
            <div style="font-size:0.75rem;color:#3e6470;">questions</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<div style='margin-top:2rem;'>", unsafe_allow_html=True)

    # Per-question results
    st.markdown("""<div style="font-size:0.7rem;letter-spacing:0.15em;text-transform:uppercase;color:#00d4b4;font-family:'DM Mono',monospace;margin-bottom:1rem;">Per Question Breakdown</div>""", unsafe_allow_html=True)

    for i, (q, ev) in enumerate(zip(questions, evaluations)):
        verdict = ev.get("verdict", "Fail")
        score = ev.get("score", 0)
        feedback = ev.get("feedback", "")
        badge_class = f"badge-{'pass' if verdict=='Pass' else 'partial' if verdict=='Partial' else 'fail'}"
        card_class = f"report-{'pass' if verdict=='Pass' else 'partial' if verdict=='Partial' else 'fail'}"

        st.markdown(f"""
        <div class="report-card {card_class}">
            <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:0.6rem;">
                <div style="font-size:0.9rem;font-weight:600;color:#e8f4f6;">Q{i+1}: {q['title']}</div>
                <div style="display:flex;align-items:center;gap:0.8rem;">
                    <span style="font-size:0.82rem;color:#4a7a86;font-family:'DM Mono',monospace;">{score}/100</span>
                    <span class="report-badge {badge_class}">{verdict}</span>
                </div>
            </div>
            <div style="font-size:0.82rem;color:#4a7a86;line-height:1.7;">{feedback}</div>
            <div style="display:flex;gap:1.5rem;margin-top:0.6rem;">
                <span style="font-size:0.72rem;color:#2e5058;font-family:'DM Mono',monospace;">Time: {ev.get('time_complexity','N/A')}</span>
                <span style="font-size:0.72rem;color:#2e5058;font-family:'DM Mono',monospace;">Space: {ev.get('space_complexity','N/A')}</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # Final AI report
    st.markdown("<div style='margin-top:2rem;'>", unsafe_allow_html=True)
    st.markdown("""<div style="font-size:0.7rem;letter-spacing:0.15em;text-transform:uppercase;color:#00d4b4;font-family:'DM Mono',monospace;margin-bottom:1rem;">AI Interview Report</div>""", unsafe_allow_html=True)

    with st.spinner("Generating final report..."):
        report = generate_final_report(config, answers, evaluations)
    st.markdown("<div style='margin:1.5rem 0;'>", unsafe_allow_html=True)
    st.text_area("AI Interview Report", report, height=420)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='margin-top:2rem;'>", unsafe_allow_html=True)
    col_a, col_b = st.columns([1, 1])
    with col_a:
        st.download_button("↓ Download Report", report, file_name="interview_report.md", mime="text/markdown")
    with col_b:
        if st.button("⬡ Start New Interview"):
            fake_loader("Initializing...", 1.5)
            st.session_state.tech_stage = "setup"
            st.session_state.tech_questions = []
            st.session_state.tech_answers = {}
            st.session_state.tech_q_index = 0
            st.rerun()

    st.markdown("</div></div></div>", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════
# PAGE: HR (placeholder)
# ════════════════════════════════════════════════════════
def show_hr():
    stage = st.session_state.hr_stage

    # ── SETUP ──
    if stage == "setup":
        st.markdown("<div style='padding:1rem 3rem;'>", unsafe_allow_html=True)
        if st.button("← Back to Home", key="hr_back"):
            fake_loader("Initializing...", 1.5)
            st.session_state.page = "home"
            st.rerun()
        st.markdown("""
        <div style="margin:0.3rem 0 1 rem 0 ;">
            <div style="font-size:0.7rem;letter-spacing:0.2em;text-transform:uppercase;color:#00d4b4;font-family:'DM Mono',monospace;margin-bottom:0.5rem;">HR Interview</div>
            <h2 style="font-size:2rem;font-weight:700;color:#e8f4f6;">Configure Your Session</h2>
            <p style="font-size:0.88rem;color:#4a7a86;margin-top:0.4rem;">Set up your mock HR interview.</p>
        </div>
        """, unsafe_allow_html=True)
        with st.form("hr_setup_form"):
            col1, col2 = st.columns([1, 1])
            with col1:
                st.markdown("**Name**")
                name = st.text_input("Name", label_visibility="collapsed")
                experience = st.selectbox("Experience Level", ["Select","Fresher", "1-3 Years", "3-5 Years", "5+ Years"])
            with col2:
                st.markdown("**Job Role**")
                role = st.text_input("Job Role", label_visibility="collapsed")
    
                
                length = st.selectbox("Interview Length", ["Select","Short (6 questions)", "Medium (8 questions)", "Long (10 questions)"])
            start = st.form_submit_button("⬡  Start Interview")
        if start:
                fake_loader("Generating...", 1.5)
                try:
                    turns = {
                        "Short (6 questions)": 6,
                        "Medium (8 questions)": 8,
                        "Long (10 questions)": 10
                    }

                    st.session_state.hr_config = {
                        "name": name.strip(),
                        "role": role.strip(),
                        "experience": experience,
                        "max_turns": turns[length]   # this may fail
                    }

                    st.session_state.hr_chat = []
                    st.session_state.hr_turn = 0
                    st.session_state.hr_start_time = time.time()
                    st.session_state.hr_report = ""
                    st.session_state.hr_stage = "interview"

                    st.rerun()

                except Exception:
                    st.warning("Please complete all fields.")
                    st.markdown("</div>", unsafe_allow_html=True)

    # ── INTERVIEW ──
    elif stage == "interview":
        config = st.session_state.hr_config
        chat = st.session_state.hr_chat
        turn = st.session_state.hr_turn
        max_turns = config["max_turns"]

        st.markdown("<div style='padding:1.5rem 3rem;'>", unsafe_allow_html=True)

        # Timer
        elapsed = time.time() - st.session_state.hr_start_time
        time_limit = max_turns * 3 * 60
        remaining = max(0, time_limit - elapsed)
        if remaining <= 0:
            st.session_state.hr_stage = "report"
            st.rerun()

        timer_color = "#e24b4a" if remaining < 120 else "#f5a623" if remaining < 300 else "#00d4b4"
        m, s = divmod(int(remaining), 60)
        st.markdown(f"""
        <div style="background:rgba(8,22,26,0.9);border:1px solid rgba(0,210,180,0.12);border-radius:10px;
        padding:0.8rem 1.5rem;display:flex;justify-content:space-between;align-items:center;margin-bottom:1.5rem;">
            <div style="font-size:0.7rem;text-transform:uppercase;letter-spacing:0.1em;color:#3e6470;font-family:'DM Mono',monospace;">
                HR Interview · {config['role'] or 'General'} · {config['experience']}
            </div>
            <div style="font-size:1.1rem;font-weight:700;color:{timer_color};font-family:'DM Mono',monospace;">
                {m:02d}:{s:02d}
            </div>
            <div style="font-size:0.7rem;color:#3e6470;font-family:'DM Mono',monospace;">
                Turn {turn} / {max_turns}
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Generate interviewer question if it's the AI's turn
        if len(chat) == 0 or chat[-1]["role"] == "candidate":
            if turn < max_turns:
                history_text = "\n".join([f"{'Interviewer' if m['role']=='interviewer' else config['name']}: {m['text']}" for m in chat])
                if turn == 0:
                    prompt = f"""You are a professional HR interviewer. You are interviewing {config['name']} for the role of {config['role'] or 'a position'}.
They have {config['experience']} of experience. Start the interview naturally and warmly. Ask them to introduce themselves. Keep it short — 2-3 sentences max."""
                elif turn == max_turns - 1:
                    prompt = f"""You are a professional HR interviewer. This is the last question. 
Conversation so far:\n{history_text}\nAsk a closing question like 'Do you have any questions for us?' Keep it natural and brief."""
                else:
                    prompt = f"""You are a professional HR interviewer interviewing {config['name']} for {config['role'] or 'a position'}.
Conversation so far:\n{history_text}\nAsk the next natural HR interview question. React briefly to their last answer (1 sentence), then ask the next question. Keep total response under 3 sentences. Be conversational, not robotic."""
                with st.spinner("Interviewer is typing..."):
                    try:
                        response = client.models.generate_content(model=MODEL_ID, contents=prompt)
                        question = response.text.strip()
                    except Exception:
                        st.warning("⚠️ System is experiencing high demand. Please restart the interview later.")
                        question= "Sorry, I’m having trouble generating a question right now. Let’s continue — tell me about yourself."
                    
                    
                chat.append({"role": "interviewer", "text": question})
                st.session_state.hr_chat = chat
                st.session_state.hr_turn = turn + 1
                st.rerun()

        # Display chat
        for msg in chat:
            if msg["role"] == "interviewer":
                st.markdown(f"""
                <div style="display:flex;gap:0.8rem;margin-bottom:1rem;align-items:flex-start;">
                    <div style="width:36px;height:36px;border-radius:50%;background:rgba(0,210,180,0.15);border:1px solid rgba(0,210,180,0.3);display:flex;align-items:center;justify-content:center;font-size:1rem;flex-shrink:0;">🎙</div>
                    <div style="background:rgba(8,22,26,0.85);border:1px solid rgba(0,210,180,0.12);border-radius:0 12px 12px 12px;padding:0.8rem 1.2rem;max-width:70%;">
                        <div style="font-size:0.65rem;color:#3e6470;font-family:'DM Mono',monospace;margin-bottom:0.3rem;">INTERVIEWER</div>
                        <div style="font-size:0.9rem;color:#e0eef2;line-height:1.7;">{msg['text']}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div style="display:flex;gap:0.8rem;margin-bottom:1rem;align-items:flex-start;flex-direction:row-reverse;">
                    <div style="width:36px;height:36px;border-radius:50%;background:rgba(0,210,180,0.1);border:1px solid rgba(0,210,180,0.2);display:flex;align-items:center;justify-content:center;font-size:1rem;flex-shrink:0;">👤</div>
                    <div style="background:rgba(0,210,180,0.06);border:1px solid rgba(0,210,180,0.15);border-radius:12px 0 12px 12px;padding:0.8rem 1.2rem;max-width:70%;">
                        <div style="font-size:0.65rem;color:#3e6470;font-family:'DM Mono',monospace;margin-bottom:0.3rem;">{config['name'].upper()}</div>
                        <div style="font-size:0.9rem;color:#e0eef2;line-height:1.7;">{msg['text']}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

        # Answer input — only if last message is from interviewer
        if chat and chat[-1]["role"] == "interviewer" and turn <= max_turns:
            with st.form(f"hr_answer_form_{turn}"):
                answer = st.text_area("Your answer", placeholder="Type your answer here...", height=120, label_visibility="collapsed")
                col_submit, col_finish = st.columns([1, 1])
                with col_submit:
                    submitted = st.form_submit_button("Send ▶", use_container_width=True)
                with col_finish:
                    finished = st.form_submit_button("Finish Interview ✓", use_container_width=True)
            if submitted and answer.strip():
                chat.append({"role": "candidate", "text": answer.strip()})
                st.session_state.hr_chat = chat
                st.rerun()
            if finished:
                fake_loader("Initializing...", 1.5)
                st.session_state.hr_stage = "report"
                st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

    # ── REPORT ──
    elif stage == "report":
        config = st.session_state.hr_config
        chat = st.session_state.hr_chat
        st.markdown("<div style='padding:2rem 3rem;'>", unsafe_allow_html=True)
        st.markdown("""
        <div style="margin-bottom:1.5rem;">
            <div style="font-size:0.7rem;letter-spacing:0.2em;text-transform:uppercase;color:#00d4b4;font-family:'DM Mono',monospace;margin-bottom:0.5rem;">Interview Complete</div>
            <h2 style="font-size:2rem;font-weight:700;color:#e8f4f6;">Your HR Report</h2>
        </div>
        """, unsafe_allow_html=True)

        if not st.session_state.hr_report:
            conversation = "\n".join([f"{'Interviewer' if m['role']=='interviewer' else config['name']}: {m['text']}" for m in chat])
            prompt = f"""You are an expert HR coach. Evaluate this mock HR interview.

Candidate: {config['name']}
Role: {config['role'] or 'General'}
Experience: {config['experience']}

Full conversation:
{conversation}

Write a detailed report in markdown with these sections:
## Overall Score (X/100)
## Communication Clarity
## Answer Depth & Use of Examples
## Confidence & Tone
## Strongest Answer
## Weakest Answer
## Key Improvements
## Final Recommendation
"""
            with st.spinner("Generating your HR report..."):
                response = client.models.generate_content(model=MODEL_ID, contents=prompt)
                st.session_state.hr_report = response.text

        st.text_area("HR Interview Report", st.session_state.hr_report, height=500)
        col_a, col_b = st.columns(2)
        with col_a:
            st.download_button("↓ Download Report", st.session_state.hr_report, file_name="hr_report.md", mime="text/markdown")
        with col_b:
            if st.button("⬡ New Interview", key="hr_new"):
                fake_loader("Initializing...", 1.5)
                st.session_state.hr_stage = "setup"
                st.session_state.hr_chat = []
                st.session_state.hr_turn = 0
                st.session_state.hr_report = ""
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
# ════════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════
# PAGE: RESUME BUILDER
# ════════════════════════════════════════════════════════
def show_resume():
    import base64, json, textwrap

    # ── Extra styles scoped to resume page ──
    st.markdown("""
    <style>
    .resume-wrap { padding: 2rem 3rem; }
    .resume-section-label {
        font-size: 0.7rem; letter-spacing: 0.2em; text-transform: uppercase;
        color: #00d4b4; font-family: 'DM Mono', monospace; margin-bottom: 0.4rem;
    }
    .resume-section-title {
        font-size: 2rem; font-weight: 700; color: #e8f4f6;
        letter-spacing: -0.02em; margin-bottom: 0.4rem;
    }
    .resume-section-sub {
        font-size: 0.85rem; color: #4a7a86; font-family: 'DM Mono', monospace;
        margin-bottom: 2rem;
    }
    /* card panels */
    .resume-card {
        background: rgba(8,22,26,0.85);
        border: 1px solid rgba(0,210,180,0.12);
        border-radius: 14px; padding: 1.8rem 2rem; margin-bottom: 1.4rem;
    }
    .resume-card-title {
        font-size: 0.72rem; letter-spacing: 0.15em; text-transform: uppercase;
        color: #00d4b4; font-family: 'DM Mono', monospace; margin-bottom: 1.2rem;
        border-bottom: 1px solid rgba(0,210,180,0.08); padding-bottom: 0.6rem;
    }
    /* text inputs inside resume form */
    [data-testid="stTextInput"] input,
    [data-testid="stTextArea"] textarea {
        background: rgba(10,26,30,0.9) !important;
        border: 1px solid rgba(0,210,180,0.2) !important;
        border-radius: 10px !important; color: #e0eef2 !important;
        font-family: 'Sora', sans-serif !important; font-size: 0.88rem !important;
    }
    [data-testid="stTextInput"] label,
    [data-testid="stTextArea"] label {
        color: #8eb8c0 !important; font-family: 'DM Mono', monospace !important;
        font-size: 0.7rem !important; letter-spacing: 0.08em !important;
        text-transform: uppercase !important; margin-bottom: 0.3rem !important;
    }
    /* format selector pills */
    .format-pills { display: flex; gap: 0.8rem; margin-bottom: 1.2rem; flex-wrap: wrap; }
    .format-pill {
        padding: 0.5rem 1.2rem; border-radius: 20px;
        border: 1px solid rgba(0,210,180,0.25); color: #4a7a86;
        font-size: 0.8rem; font-family: 'DM Mono', monospace; cursor: pointer;
        background: rgba(8,22,26,0.6); transition: all 0.2s;
    }
    .format-pill.active {
        border-color: #00d4b4; color: #00d4b4;
        background: rgba(0,210,180,0.08);
        box-shadow: 0 0 12px rgba(0,210,180,0.1);
    }
    /* preview box */
    .resume-preview-box {
        background: #ffffff; border-radius: 10px;
        padding: 2.5rem; color: #1a1a1a; font-family: Arial, sans-serif;
        font-size: 0.88rem; line-height: 1.6; min-height: 500px;
        box-shadow: 0 0 40px rgba(0,0,0,0.4);
    }
    /* submit button override — keep it inline, not full width */
    div[data-testid="stFormSubmitButton"] {
        display: flex !important;
        justify-content: flex-start !important;
    }
    div[data-testid="stFormSubmitButton"] button {
        width: auto !important;
        min-width: 180px !important;
        padding: 0.6rem 2rem !important;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("<div class='resume-wrap'>", unsafe_allow_html=True)

    # Back button
    if st.button("← Back to Home", key="resume_back"):
        fake_loader("Going Back...", 1.5)
        st.session_state.page = "home"
        st.rerun()

    st.markdown("""
    <div style="margin: 1rem 0 2rem 0;">
        <div class="resume-section-label">AI-Powered</div>
        <div class="resume-section-title">Resume Builder</div>
        <div class="resume-section-sub">Fill in your details · choose a format · preview & download</div>
    </div>
    """, unsafe_allow_html=True)

    # ── SESSION STATE for resume data ──
    if "resume_data" not in st.session_state:
        st.session_state.resume_data = {}
    if "resume_format" not in st.session_state:
        st.session_state.resume_format = "Classic"
    if "resume_preview_html" not in st.session_state:
        st.session_state.resume_preview_html = ""
    if "resume_stage" not in st.session_state:
        st.session_state.resume_stage = "form"   # form | preview

    # ════════════════════════════════════
    # STAGE: FORM
    # ════════════════════════════════════
    if st.session_state.resume_stage == "form":

        with st.form("resume_form", clear_on_submit=False):

            # ── Personal Info ──
            st.markdown("<div class='resume-card'><div class='resume-card-title'>◈ Personal Information</div>", unsafe_allow_html=True)
            col1, col2 = st.columns(2)
            with col1:
                full_name   = st.text_input("Full Name", placeholder="Full Name")
                email       = st.text_input("Email Address", placeholder="Email Address")
                phone       = st.text_input("Phone Number", placeholder="Phone Number")
            with col2:
                job_title   = st.text_input("Desired Job Title", placeholder="Desired Job Title")
                location    = st.text_input("Location", placeholder="City, Country")
                linkedin    = st.text_input("LinkedIn / Portfolio URL", placeholder="LinkedIn or Portfolio URL")
            st.markdown("</div>", unsafe_allow_html=True)

            # ── Summary ──
            st.markdown("<div class='resume-card'><div class='resume-card-title'>◈ Professional Summary</div>", unsafe_allow_html=True)
            summary = st.text_area("Brief summary about yourself", height=100,
                placeholder="Professional Summary")
            st.markdown("</div>", unsafe_allow_html=True)

            # ── Experience ──
            st.markdown("<div class='resume-card'><div class='resume-card-title'>◈ Work Experience</div>", unsafe_allow_html=True)
            col_e1, col_e2 = st.columns(2)
            with col_e1:
                exp1_title   = st.text_input("Job Title #1", placeholder="Job Title")
                exp1_company = st.text_input("Company #1", placeholder="Company Name")
                exp1_dates   = st.text_input("Duration #1", placeholder="Start Date – End Date")
            with col_e2:
                exp2_title   = st.text_input("Job Title #2 (optional)", placeholder="Job Title (optional)")
                exp2_company = st.text_input("Company #2 (optional)", placeholder="Company Name (optional)")
                exp2_dates   = st.text_input("Duration #2 (optional)", placeholder="Start Date – End Date")
            exp1_desc = st.text_area("Responsibilities / Achievements #1", height=80,
                placeholder="Responsibilities & Achievements")
            exp2_desc = st.text_area("Responsibilities / Achievements #2 (optional)", height=80,
                placeholder="Responsibilities & Achievements (optional)")
            st.markdown("</div>", unsafe_allow_html=True)

            # ── Education ──
            st.markdown("<div class='resume-card'><div class='resume-card-title'>◈ Education</div>", unsafe_allow_html=True)
            col_ed1, col_ed2 = st.columns(2)
            with col_ed1:
                edu_degree  = st.text_input("Degree", placeholder="Degree / Field of Study")
                edu_college = st.text_input("Institution", placeholder="College / University")
            with col_ed2:
                edu_year    = st.text_input("Year of Graduation", placeholder="Graduation Year")
                edu_grade   = st.text_input("Grade / CGPA", placeholder="Grade / CGPA")
            st.markdown("</div>", unsafe_allow_html=True)

            # ── Skills ──
            st.markdown("<div class='resume-card'><div class='resume-card-title'>◈ Skills</div>", unsafe_allow_html=True)
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                tech_skills = st.text_input("Technical Skills", placeholder="Technical Skills")
            with col_s2:
                soft_skills = st.text_input("Soft Skills", placeholder="Soft Skills")
            st.markdown("</div>", unsafe_allow_html=True)

            # ── Projects ──
            st.markdown("<div class='resume-card'><div class='resume-card-title'>◈ Projects</div>", unsafe_allow_html=True)
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                proj1_name = st.text_input("Project Name #1", placeholder="Project Name")
                proj1_tech = st.text_input("Tech Stack #1", placeholder="Technologies Used")
            with col_p2:
                proj2_name = st.text_input("Project Name #2 (optional)", placeholder="Project Name (optional)")
                proj2_tech = st.text_input("Tech Stack #2 (optional)", placeholder="Technologies Used (optional)")
            proj1_desc = st.text_area("Description #1", height=70, placeholder="Project Description")
            proj2_desc = st.text_area("Description #2 (optional)", height=70, placeholder="Project Description (optional)")
            st.markdown("</div>", unsafe_allow_html=True)

            # ── Certifications & Achievements ──
            st.markdown("<div class='resume-card'><div class='resume-card-title'>◈ Certifications & Achievements</div>", unsafe_allow_html=True)
            certifications = st.text_area("List certifications / achievements (one per line)", height=80,
                placeholder="Certifications & Achievements (one per line)")
            st.markdown("</div>", unsafe_allow_html=True)

            # ── Format selector ──
            st.markdown("<div class='resume-card'><div class='resume-card-title'>◈ Choose Resume Format</div>", unsafe_allow_html=True)
            resume_format = st.selectbox(
                "Format",
                ["Classic", "Modern", "Minimal"],
                index=["Classic", "Modern", "Minimal"].index(st.session_state.resume_format),
                help="Classic = traditional black & white | Modern = teal accent sidebar | Minimal = clean single column"
            )
            st.markdown("""
            <div style="display:flex;gap:1.5rem;margin-top:0.8rem;flex-wrap:wrap;">
                <div style="font-size:0.78rem;color:#4a7a86;">
                    <span style="color:#00d4b4;font-family:'DM Mono',monospace;">Classic</span> — Traditional, ATS-friendly, black & white
                </div>
                <div style="font-size:0.78rem;color:#4a7a86;">
                    <span style="color:#00d4b4;font-family:'DM Mono',monospace;">Modern</span> — Two-column, teal accent sidebar
                </div>
                <div style="font-size:0.78rem;color:#4a7a86;">
                    <span style="color:#00d4b4;font-family:'DM Mono',monospace;">Minimal</span> — Clean single-column, generous spacing
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

            # ── SUBMIT (inline, not full width — controlled by CSS above) ──
            submitted = st.form_submit_button("⬡  Generate Preview")

        # ── Handle submission ──
        if submitted:
            if not full_name or not email or not job_title:
                st.warning("⚠️ Please fill in at least Name, Email, and Job Title.")
            else:
                st.session_state.resume_data = {
                    "full_name": full_name, "email": email, "phone": phone,
                    "job_title": job_title, "location": location, "linkedin": linkedin,
                    "summary": summary,
                    "exp1_title": exp1_title, "exp1_company": exp1_company,
                    "exp1_dates": exp1_dates, "exp1_desc": exp1_desc,
                    "exp2_title": exp2_title, "exp2_company": exp2_company,
                    "exp2_dates": exp2_dates, "exp2_desc": exp2_desc,
                    "edu_degree": edu_degree, "edu_college": edu_college,
                    "edu_year": edu_year, "edu_grade": edu_grade,
                    "tech_skills": tech_skills, "soft_skills": soft_skills,
                    "proj1_name": proj1_name, "proj1_tech": proj1_tech, "proj1_desc": proj1_desc,
                    "proj2_name": proj2_name, "proj2_tech": proj2_tech, "proj2_desc": proj2_desc,
                    "certifications": certifications,
                    "format": resume_format,
                }
                st.session_state.resume_format = resume_format
                fake_loader("Building your resume...", 1.5)
                st.session_state.resume_stage = "preview"
                st.rerun()

    # ════════════════════════════════════
    # STAGE: PREVIEW + DOWNLOAD
    # ════════════════════════════════════
    elif st.session_state.resume_stage == "preview":
        d = st.session_state.resume_data
        fmt = d.get("format", "Classic")

        # ── Build HTML preview ──
        def nl2bullets(text):
            if not text: return ""
            lines = [l.strip().lstrip("•-").strip() for l in text.strip().split("\n") if l.strip()]
            return "".join(f"<li>{l}</li>" for l in lines)

        def safe(val): return val if val else ""

        if fmt == "Classic":
            
            # Build sections separately to avoid f-string quote conflicts
            contact_parts = []
            if safe(d['email']): contact_parts.append(safe(d['email']))
            if safe(d['phone']): contact_parts.append(safe(d['phone']))
            if safe(d['location']): contact_parts.append(safe(d['location']))
            if safe(d['linkedin']): contact_parts.append(safe(d['linkedin']))
            contact_line = ' · '.join(contact_parts)

            summary_html = f'<div><h2 style="font-size:13px;text-transform:uppercase;letter-spacing:1px;border-bottom:1px solid #999;padding-bottom:3px;margin-bottom:8px;">Summary</h2><p style="margin:0 0 14px;">{safe(d["summary"])}</p></div>' if d.get('summary') else ''

            exp_html = ''
            if d.get('exp1_title'):
                exp_html += '<h2 style="font-size:13px;text-transform:uppercase;letter-spacing:1px;border-bottom:1px solid #999;padding-bottom:3px;margin-bottom:8px;">Work Experience</h2>'
                exp_html += f'<div style="margin-bottom:10px;"><div style="display:flex;justify-content:space-between;"><strong>{safe(d["exp1_title"])} — {safe(d["exp1_company"])}</strong><span style="font-size:12px;color:#555;">{safe(d["exp1_dates"])}</span></div><ul style="margin:4px 0 0 16px;padding:0;">{nl2bullets(d.get("exp1_desc",""))}</ul></div>'
                if d.get('exp2_title'):
                    exp_html += f'<div style="margin-bottom:10px;"><div style="display:flex;justify-content:space-between;"><strong>{safe(d["exp2_title"])} — {safe(d["exp2_company"])}</strong><span style="font-size:12px;color:#555;">{safe(d["exp2_dates"])}</span></div><ul style="margin:4px 0 0 16px;padding:0;">{nl2bullets(d.get("exp2_desc",""))}</ul></div>'
                exp_html = f'<div>{exp_html}</div>'

            edu_html = ''
            if d.get('edu_degree'):
                edu_html = f'<div><h2 style="font-size:13px;text-transform:uppercase;letter-spacing:1px;border-bottom:1px solid #999;padding-bottom:3px;margin-bottom:8px;">Education</h2><div style="display:flex;justify-content:space-between;"><span><strong>{safe(d["edu_degree"])}</strong> — {safe(d["edu_college"])}</span><span style="font-size:12px;color:#555;">{safe(d["edu_year"])} | {safe(d["edu_grade"])}</span></div></div><br/>'

            skills_html = ''
            if d.get('tech_skills'):
                skills_html = f'<div><h2 style="font-size:13px;text-transform:uppercase;letter-spacing:1px;border-bottom:1px solid #999;padding-bottom:3px;margin-bottom:8px;">Skills</h2><p style="margin:0 0 14px;"><strong>Technical:</strong> {safe(d["tech_skills"])}<br/><strong>Soft:</strong> {safe(d["soft_skills"])}</p></div>'

            proj_html = ''
            if d.get('proj1_name'):
                proj_html = '<h2 style="font-size:13px;text-transform:uppercase;letter-spacing:1px;border-bottom:1px solid #999;padding-bottom:3px;margin-bottom:8px;">Projects</h2>'
                proj_html += f'<div style="margin-bottom:8px;"><strong>{safe(d["proj1_name"])}</strong> <span style="color:#555;font-size:12px;">({safe(d["proj1_tech"])})</span><br/>{safe(d["proj1_desc"])}</div>'
                if d.get('proj2_name'):
                    proj_html += f'<div style="margin-bottom:8px;"><strong>{safe(d["proj2_name"])}</strong> <span style="color:#555;font-size:12px;">({safe(d["proj2_tech"])})</span><br/>{safe(d["proj2_desc"])}</div>'
                proj_html = f'<div>{proj_html}</div>'

            cert_html = ''
            if d.get('certifications'):
                cert_html = f'<div><h2 style="font-size:13px;text-transform:uppercase;letter-spacing:1px;border-bottom:1px solid #999;padding-bottom:3px;margin-bottom:8px;">Certifications & Achievements</h2><ul style="margin:0 0 0 16px;padding:0;">{nl2bullets(d["certifications"])}</ul></div>'

            preview_html = f"""
            <div style="font-family:Arial,sans-serif;max-width:750px;margin:0 auto;color:#1a1a1a;font-size:14px;line-height:1.6;">
                <div style="text-align:center;border-bottom:2px solid #000;padding-bottom:12px;margin-bottom:16px;">
                    <h1 style="margin:0;font-size:24px;letter-spacing:1px;">{safe(d['full_name']).upper()}</h1>
                    <div style="font-size:13px;color:#333;margin-top:4px;">{safe(d['job_title'])}</div>
                    <div style="font-size:12px;color:#555;margin-top:4px;">{contact_line}</div>
                </div>
                {summary_html}{exp_html}{edu_html}{skills_html}{proj_html}{cert_html}
            </div>"""

        elif fmt == "Modern":
            preview_html = f"""
            <div style="font-family:Arial,sans-serif;max-width:750px;margin:0 auto;display:flex;color:#1a1a1a;font-size:13px;line-height:1.6;">
                <div style="width:220px;min-width:220px;background:#006b5e;color:#fff;padding:24px 16px;border-radius:8px 0 0 8px;">
                    <h1 style="margin:0 0 4px;font-size:18px;line-height:1.2;">{safe(d['full_name'])}</h1>
                    <div style="font-size:12px;opacity:0.85;margin-bottom:16px;">{safe(d['job_title'])}</div>
                    <div style="font-size:11px;opacity:0.8;margin-bottom:4px;">📧 {safe(d['email'])}</div>
                    <div style="font-size:11px;opacity:0.8;margin-bottom:4px;">📞 {safe(d['phone'])}</div>
                    <div style="font-size:11px;opacity:0.8;margin-bottom:4px;">📍 {safe(d['location'])}</div>
                    <div style="font-size:11px;opacity:0.8;margin-bottom:16px;">🔗 {safe(d['linkedin'])}</div>
                    {'<div style="border-top:1px solid rgba(255,255,255,0.3);padding-top:12px;margin-bottom:8px;font-size:11px;font-weight:700;letter-spacing:1px;text-transform:uppercase;">Skills</div><div style="font-size:11px;opacity:0.9;margin-bottom:6px;"><strong>Technical</strong><br/>' + safe(d['tech_skills']) + '</div><div style="font-size:11px;opacity:0.9;"><strong>Soft</strong><br/>' + safe(d['soft_skills']) + '</div>' if d['tech_skills'] else ''}
                    {'<div style="border-top:1px solid rgba(255,255,255,0.3);margin-top:12px;padding-top:12px;"><div style="font-size:11px;font-weight:700;letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;">Education</div><div style="font-size:11px;opacity:0.9;"><strong>' + safe(d['edu_degree']) + '</strong><br/>' + safe(d['edu_college']) + '<br/>' + safe(d['edu_year']) + ' · ' + safe(d['edu_grade']) + '</div></div>' if d['edu_degree'] else ''}
                </div>
                <div style="flex:1;padding:24px 20px;background:#fff;border-radius:0 8px 8px 0;border:1px solid #e0e0e0;border-left:none;">
                    {'<p style="margin:0 0 16px;color:#444;">' + safe(d['summary']) + '</p>' if d['summary'] else ''}
                    {'<h2 style="font-size:13px;color:#006b5e;text-transform:uppercase;letter-spacing:1px;border-bottom:2px solid #006b5e;padding-bottom:3px;margin-bottom:10px;">Experience</h2>' if d['exp1_title'] else ''}
                    {'<div style="margin-bottom:10px;"><div style="display:flex;justify-content:space-between;"><strong>' + safe(d['exp1_title']) + '</strong><span style="font-size:11px;color:#777;">' + safe(d['exp1_dates']) + '</span></div><div style="color:#006b5e;font-size:12px;">' + safe(d['exp1_company']) + '</div><ul style="margin:4px 0 0 14px;padding:0;font-size:12px;">' + nl2bullets(d['exp1_desc']) + '</ul></div>' if d['exp1_title'] else ''}
                    {'<div style="margin-bottom:10px;"><div style="display:flex;justify-content:space-between;"><strong>' + safe(d['exp2_title']) + '</strong><span style="font-size:11px;color:#777;">' + safe(d['exp2_dates']) + '</span></div><div style="color:#006b5e;font-size:12px;">' + safe(d['exp2_company']) + '</div><ul style="margin:4px 0 0 14px;padding:0;font-size:12px;">' + nl2bullets(d['exp2_desc']) + '</ul></div>' if d['exp2_title'] else ''}
                    {'<h2 style="font-size:13px;color:#006b5e;text-transform:uppercase;letter-spacing:1px;border-bottom:2px solid #006b5e;padding-bottom:3px;margin-bottom:10px;margin-top:14px;">Projects</h2>' if d['proj1_name'] else ''}
                    {'<div style="margin-bottom:8px;"><strong>' + safe(d['proj1_name']) + '</strong> <span style="color:#777;font-size:11px;">· ' + safe(d['proj1_tech']) + '</span><br/><span style="font-size:12px;">' + safe(d['proj1_desc']) + '</span></div>' if d['proj1_name'] else ''}
                    {'<div style="margin-bottom:8px;"><strong>' + safe(d['proj2_name']) + '</strong> <span style="color:#777;font-size:11px;">· ' + safe(d['proj2_tech']) + '</span><br/><span style="font-size:12px;">' + safe(d['proj2_desc']) + '</span></div>' if d['proj2_name'] else ''}
                    {'<h2 style="font-size:13px;color:#006b5e;text-transform:uppercase;letter-spacing:1px;border-bottom:2px solid #006b5e;padding-bottom:3px;margin-bottom:10px;margin-top:14px;">Certifications</h2><ul style="margin:0 0 0 14px;padding:0;font-size:12px;">' + nl2bullets(d['certifications']) + '</ul>' if d['certifications'] else ''}
                </div>
            </div>"""

        else:  # Minimal
            preview_html = f"""
            <div style="font-family:'Georgia',serif;max-width:700px;margin:0 auto;color:#1a1a1a;font-size:14px;line-height:1.8;">
                <h1 style="margin:0;font-size:28px;font-weight:normal;letter-spacing:2px;">{safe(d['full_name'])}</h1>
                <div style="font-size:14px;color:#555;margin:4px 0 6px;">{safe(d['job_title'])}</div>
                <div style="font-size:12px;color:#888;margin-bottom:24px;">
                    {safe(d['email'])} &nbsp;|&nbsp; {safe(d['phone'])} &nbsp;|&nbsp; {safe(d['location'])} &nbsp;|&nbsp; {safe(d['linkedin'])}
                </div>
                {'<p style="margin:0 0 24px;color:#333;">' + safe(d['summary']) + '</p>' if d['summary'] else ''}
                {'<h2 style="font-size:11px;text-transform:uppercase;letter-spacing:2px;color:#888;margin-bottom:8px;">Experience</h2><hr style="border:none;border-top:1px solid #ddd;margin-bottom:14px;"/>' if d['exp1_title'] else ''}
                {'<div style="margin-bottom:14px;"><div style="display:flex;justify-content:space-between;font-size:13px;"><span><em>' + safe(d['exp1_title']) + '</em>, ' + safe(d['exp1_company']) + '</span><span style="color:#888;">' + safe(d['exp1_dates']) + '</span></div><ul style="margin:4px 0 0 16px;padding:0;font-size:13px;">' + nl2bullets(d['exp1_desc']) + '</ul></div>' if d['exp1_title'] else ''}
                {'<div style="margin-bottom:14px;"><div style="display:flex;justify-content:space-between;font-size:13px;"><span><em>' + safe(d['exp2_title']) + '</em>, ' + safe(d['exp2_company']) + '</span><span style="color:#888;">' + safe(d['exp2_dates']) + '</span></div><ul style="margin:4px 0 0 16px;padding:0;font-size:13px;">' + nl2bullets(d['exp2_desc']) + '</ul></div>' if d['exp2_title'] else ''}
                {'<h2 style="font-size:11px;text-transform:uppercase;letter-spacing:2px;color:#888;margin-bottom:8px;margin-top:20px;">Education</h2><hr style="border:none;border-top:1px solid #ddd;margin-bottom:14px;"/><p style="margin:0 0 20px;font-size:13px;"><em>' + safe(d['edu_degree']) + '</em>, ' + safe(d['edu_college']) + ' · ' + safe(d['edu_year']) + ' · ' + safe(d['edu_grade']) + '</p>' if d['edu_degree'] else ''}
                {'<h2 style="font-size:11px;text-transform:uppercase;letter-spacing:2px;color:#888;margin-bottom:8px;">Skills</h2><hr style="border:none;border-top:1px solid #ddd;margin-bottom:10px;"/><p style="margin:0 0 20px;font-size:13px;">' + safe(d['tech_skills']) + ' &nbsp;·&nbsp; ' + safe(d['soft_skills']) + '</p>' if d['tech_skills'] else ''}
                {'<h2 style="font-size:11px;text-transform:uppercase;letter-spacing:2px;color:#888;margin-bottom:8px;">Projects</h2><hr style="border:none;border-top:1px solid #ddd;margin-bottom:10px;"/>' if d['proj1_name'] else ''}
                {'<div style="margin-bottom:10px;font-size:13px;"><strong>' + safe(d['proj1_name']) + '</strong> <span style="color:#888;">(' + safe(d['proj1_tech']) + ')</span><br/>' + safe(d['proj1_desc']) + '</div>' if d['proj1_name'] else ''}
                {'<div style="margin-bottom:10px;font-size:13px;"><strong>' + safe(d['proj2_name']) + '</strong> <span style="color:#888;">(' + safe(d['proj2_tech']) + ')</span><br/>' + safe(d['proj2_desc']) + '</div>' if d['proj2_name'] else ''}
                {'<h2 style="font-size:11px;text-transform:uppercase;letter-spacing:2px;color:#888;margin-bottom:8px;margin-top:16px;">Certifications</h2><hr style="border:none;border-top:1px solid #ddd;margin-bottom:10px;"/><ul style="margin:0 0 0 16px;padding:0;font-size:13px;">' + nl2bullets(d['certifications']) + '</ul>' if d['certifications'] else ''}
            </div>"""

        # ── Show preview ──
        st.markdown("""
        <div style="margin-bottom:1rem;">
            <div class="resume-section-label">Preview</div>
            <div style="font-size:1.2rem;font-weight:600;color:#e8f4f6;margin-bottom:0.3rem;">
                Your Resume — <span style="color:#00d4b4;">""" + fmt + """ Format</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown(
            '<div class="resume-preview-box">' + preview_html + '</div>',
            unsafe_allow_html=True
        )

        st.markdown("<div style='height:1.5rem;'></div>", unsafe_allow_html=True)

        # ── Download buttons ──
        col_back, col_word, col_pdf, col_spacer = st.columns([1, 1, 1, 2])

        with col_back:
            if st.button("← Edit Resume", key="resume_edit"):
                st.session_state.resume_stage = "form"
                st.rerun()

        # ── Shared DOCX builder ──
        def build_docx_bytes(d, safe):
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import io

            doc = DocxDocument()
            for sec in doc.sections:
                sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Pt(72)

            def add_heading(doc, title):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(title.upper())
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '000000')
                pBdr.append(bottom)
                pPr.append(pBdr)

            def add_bullet(doc, text):
                if not text: return
                for line in text.strip().split('\n'):
                    line = line.strip().lstrip('•-').strip()
                    if line:
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.space_after = Pt(2)
                        run = p.add_run(line)
                        run.font.size = Pt(10)

            # Name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(safe(d['full_name']))
            run.bold = True; run.font.size = Pt(18)

            # Job title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(safe(d['job_title']))
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(68, 68, 68)

            # Contact
            contact_line = '  ·  '.join([x for x in [safe(d['email']), safe(d['phone']), safe(d['location']), safe(d['linkedin'])] if x])
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(contact_line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(85, 85, 85)

            # Summary
            if d.get('summary'):
                add_heading(doc, 'Summary')
                p = doc.add_paragraph(safe(d['summary']))
                p.paragraph_format.space_after = Pt(6)
                p.runs[0].font.size = Pt(10)

            # Experience
            if d.get('exp1_title'):
                add_heading(doc, 'Work Experience')
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r1 = p.add_run(f"{safe(d['exp1_title'])} — {safe(d['exp1_company'])}")
                r1.bold = True; r1.font.size = Pt(10)
                r2 = p.add_run(f"   {safe(d['exp1_dates'])}")
                r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(119,119,119)
                add_bullet(doc, d.get('exp1_desc',''))
            if d.get('exp2_title'):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
                r1 = p.add_run(f"{safe(d['exp2_title'])} — {safe(d['exp2_company'])}")
                r1.bold = True; r1.font.size = Pt(10)
                r2 = p.add_run(f"   {safe(d['exp2_dates'])}")
                r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(119,119,119)
                add_bullet(doc, d.get('exp2_desc',''))

            # Education
            if d.get('edu_degree'):
                add_heading(doc, 'Education')
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                r1 = p.add_run(f"{safe(d['edu_degree'])} — {safe(d['edu_college'])}")
                r1.bold = True; r1.font.size = Pt(10)
                r2 = p.add_run(f"   {safe(d['edu_year'])}  |  {safe(d['edu_grade'])}")
                r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(119,119,119)

            # Skills
            if d.get('tech_skills'):
                add_heading(doc, 'Skills')
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r1 = p.add_run('Technical: '); r1.bold = True; r1.font.size = Pt(10)
                r2 = p.add_run(safe(d['tech_skills'])); r2.font.size = Pt(10)
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(6)
                r3 = p2.add_run('Soft: '); r3.bold = True; r3.font.size = Pt(10)
                r4 = p2.add_run(safe(d['soft_skills'])); r4.font.size = Pt(10)

            # Projects
            if d.get('proj1_name'):
                add_heading(doc, 'Projects')
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r1 = p.add_run(safe(d['proj1_name'])); r1.bold = True; r1.font.size = Pt(10)
                r2 = p.add_run(f"  ({safe(d['proj1_tech'])})"); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(85,85,85)
                if d.get('proj1_desc'):
                    p2 = doc.add_paragraph(safe(d['proj1_desc']))
                    p2.paragraph_format.space_after = Pt(4); p2.runs[0].font.size = Pt(10)
            if d.get('proj2_name'):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r1 = p.add_run(safe(d['proj2_name'])); r1.bold = True; r1.font.size = Pt(10)
                r2 = p.add_run(f"  ({safe(d['proj2_tech'])})"); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(85,85,85)
                if d.get('proj2_desc'):
                    p2 = doc.add_paragraph(safe(d['proj2_desc']))
                    p2.paragraph_format.space_after = Pt(4); p2.runs[0].font.size = Pt(10)

            # Certifications
            if d.get('certifications'):
                add_heading(doc, 'Certifications & Achievements')
                add_bullet(doc, d['certifications'])

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        # ── Word DOCX button ──
        with col_word:
            if st.button("↓ Download Word", key="dl_docx"):
                try:
                    buf = build_docx_bytes(d, safe)
                    st.download_button(
                        "⬇ Click to save .docx",
                        data=buf.getvalue(),
                        file_name=f"{d['full_name'].replace(' ','_')}_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_docx_final"
                    )
                except ImportError:
                    st.warning("⚠️ Run: pip install python-docx")
                except Exception as e:
                    st.error(f"⚠️ Error: {e}")

        # ── PDF button ──
        with col_pdf:
            if st.button("↓ Download PDF", key="dl_pdf"):
                try:
                    import subprocess, os, tempfile
                    buf = build_docx_bytes(d, safe)

                    # Write docx to temp file
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(buf.getvalue())
                        tmp_path = tmp.name

                    tmp_dir = os.path.dirname(tmp_path)

                    # Convert with LibreOffice
                    result = subprocess.run(
                        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp_dir, tmp_path],
                        capture_output=True, text=True, timeout=30
                    )

                    pdf_path = tmp_path.replace(".docx", ".pdf")

                    if os.path.exists(pdf_path):
                        with open(pdf_path, "rb") as f:
                            pdf_bytes = f.read()
                        st.download_button(
                            "⬇ Click to save .pdf",
                            data=pdf_bytes,
                            file_name=f"{d['full_name'].replace(' ','_')}_resume.pdf",
                            mime="application/pdf",
                            key="dl_pdf_final"
                        )
                        os.unlink(tmp_path)
                        os.unlink(pdf_path)
                    else:
                        st.warning("⚠️ LibreOffice not found. Install it from libreoffice.org, then restart.")

                except FileNotFoundError:
                    st.warning("⚠️ LibreOffice not installed. Download from libreoffice.org to enable PDF export.")
                except ImportError:
                    st.warning("⚠️ Run: pip install python-docx")
                except Exception as e:
                    st.error(f"⚠️ PDF error: {e}")

    st.markdown("</div>", unsafe_allow_html=True)
    # ── Resume session state (add with the others at top, or here works too) ──
if "resume_data" not in st.session_state:
    st.session_state.resume_data = {}
if "resume_format" not in st.session_state:
    st.session_state.resume_format = "Classic"
if "resume_stage" not in st.session_state:
    st.session_state.resume_stage = "form"

# ════════════════════════════════════════════════════════
# ROUTER
# ════════════════════════════════════════════════════════
page = st.session_state.page

if page == "home":
    show_home()
elif page == "technical":
    stage = st.session_state.tech_stage
    if stage == "setup":
        show_technical_setup()
    elif stage == "interview":
        show_technical_interview()
    elif stage == "report":
        show_technical_report()
elif page == "hr":
    show_hr()
elif page == "resume":
    show_resume()
